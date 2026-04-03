import io
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_pdf(conversation, messages) -> bytes:
    """Génère un rapport PDF d'une conversation."""
    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(
        buffer,
        pagesize    = A4,
        leftMargin  = 2 * cm,
        rightMargin = 2 * cm,
        topMargin   = 2 * cm,
        bottomMargin= 2 * cm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "Title",
        parent    = styles["Normal"],
        fontSize  = 20,
        textColor = colors.HexColor("#6366f1"),
        spaceAfter= 4,
        alignment = TA_CENTER,
        fontName  = "Helvetica-Bold",
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent    = styles["Normal"],
        fontSize  = 10,
        textColor = colors.HexColor("#94a3b8"),
        spaceAfter= 16,
        alignment = TA_CENTER,
    )
    user_label_style = ParagraphStyle(
        "UserLabel",
        parent    = styles["Normal"],
        fontSize  = 9,
        textColor = colors.HexColor("#6366f1"),
        fontName  = "Helvetica-Bold",
        spaceBefore = 10,
        spaceAfter= 2,
    )
    bot_label_style = ParagraphStyle(
        "BotLabel",
        parent    = styles["Normal"],
        fontSize  = 9,
        textColor = colors.HexColor("#10b981"),
        fontName  = "Helvetica-Bold",
        spaceBefore = 10,
        spaceAfter= 2,
    )
    msg_style = ParagraphStyle(
        "Msg",
        parent    = styles["Normal"],
        fontSize  = 10,
        textColor = colors.HexColor("#1e293b"),
        leading   = 14,
        spaceAfter= 4,
    )
    footer_style = ParagraphStyle(
        "Footer",
        parent    = styles["Normal"],
        fontSize  = 8,
        textColor = colors.HexColor("#94a3b8"),
        alignment = TA_CENTER,
        spaceBefore = 20,
    )

    bot_name  = conversation.bot.name
    date_str  = conversation.created_at.strftime("%d/%m/%Y à %H:%M")
    msg_count = messages.count()

    story = [
        Paragraph("ChatFlow", title_style),
        Paragraph(f"Rapport de conversation — {bot_name}", subtitle_style),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")),
        Spacer(1, 0.3 * cm),

        # Metadata table
        Table(
            [
                ["Bot :", bot_name],
                ["Date :", date_str],
                ["Messages :", str(msg_count)],
                ["Session :", str(conversation.session_id)[:8].upper()],
            ],
            colWidths=[3 * cm, 14 * cm],
            style=TableStyle([
                ("FONTNAME",  (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE",  (0, 0), (-1, -1), 9),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#64748b")),
                ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#1e293b")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]),
        ),
        Spacer(1, 0.4 * cm),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")),
        Spacer(1, 0.3 * cm),
    ]

    for msg in messages:
        label   = "👤 Utilisateur" if msg.role == "user" else f"🤖 {bot_name}"
        l_style = user_label_style if msg.role == "user" else bot_label_style
        time_str = msg.created_at.strftime("%H:%M")

        story.append(Paragraph(f"{label}  <font size='8' color='#94a3b8'>— {time_str}</font>", l_style))
        # Escape special chars for reportlab
        content = msg.content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(content, msg_style))

    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")))
    story.append(Paragraph(
        f"Généré par ChatFlow le {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
        footer_style
    ))

    doc.build(story)
    return buffer.getvalue()


def generate_docx(conversation, messages) -> bytes:
    """Génère un rapport Word (.docx) d'une conversation."""
    doc = DocxDocument()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    bot_name = conversation.bot.name
    date_str = conversation.created_at.strftime("%d/%m/%Y à %H:%M")

    # Titre
    title      = doc.add_heading("ChatFlow", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

    subtitle = doc.add_paragraph(f"Rapport de conversation — {bot_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()

    # Metadata
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    rows_data = [
        ("Bot", bot_name),
        ("Date", date_str),
        ("Messages", str(messages.count())),
        ("Session", str(conversation.session_id)[:8].upper()),
    ]
    for i, (key, val) in enumerate(rows_data):
        meta.rows[i].cells[0].text = key
        meta.rows[i].cells[1].text = val
        meta.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # Messages
    for msg in messages:
        is_user  = msg.role == "user"
        label    = "👤 Utilisateur" if is_user else f"🤖 {bot_name}"
        time_str = msg.created_at.strftime("%H:%M")

        label_para = doc.add_paragraph()
        label_run  = label_para.add_run(f"{label} — {time_str}")
        label_run.font.bold       = True
        label_run.font.size       = Pt(9)
        label_run.font.color.rgb  = RGBColor(0x63, 0x66, 0xF1) if is_user else RGBColor(0x10, 0xB9, 0x81)

        msg_para = doc.add_paragraph(msg.content)
        msg_para.paragraph_format.left_indent = Cm(0.5)
        for run in msg_para.runs:
            run.font.size = Pt(10)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Généré par ChatFlow le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
